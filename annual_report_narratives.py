
from docx import Document
from bcompiler.utils import project_data_from_master
from collections import OrderedDict
import datetime
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor
import difflib


def converting_RAGs(rag):
    if rag == 'Green':
        return 'G'
    elif rag == 'Amber/Green':
        return 'A/G'
    elif rag == 'Amber':
        return 'A'
    elif rag == 'Amber/Red':
        return 'A/R'
    else:
        return 'R'


def cell_colouring(cell, colour):
    if colour == 'R':
        colour = parse_xml(r'<w:shd {} w:fill="cb1f00"/>'.format(nsdecls('w')))
    elif colour == 'A/R':
        colour = parse_xml(r'<w:shd {} w:fill="f97b31"/>'.format(nsdecls('w')))
    elif colour == 'A':
        colour = parse_xml(r'<w:shd {} w:fill="fce553"/>'.format(nsdecls('w')))
    elif colour == 'A/G':
        colour = parse_xml(r'<w:shd {} w:fill="a5b700"/>'.format(nsdecls('w')))
    elif colour == 'G':
        colour = parse_xml(r'<w:shd {} w:fill="17960c"/>'.format(nsdecls('w')))

    cell._tc.get_or_add_tcPr().append(colour)


'''function places text into doc highlighing all changes'''


def compare_text_showall(text_1, text_2, doc):
    comp = difflib.Differ()
    diff = list(comp.compare(text_2.split(), text_1.split()))
    new_text = diff
    y = doc.add_paragraph()

    for i in range(0, len(diff)):
        f = len(diff) - 1
        if i < f:
            a = i - 1
        else:
            a = i

        if diff[i][0:3] == '  |':
            j = i + 1
            if diff[i][0:3] and diff[a][0:3] == '  |':
                y = doc.add_paragraph()
            else:
                pass
        elif diff[i][0:3] == '+ |':
            if diff[i][0:3] and diff[a][0:3] == '+ |':
                y = doc.add_paragraph()
            else:
                pass
        elif diff[i][0:3] == '- |':
            pass
        elif diff[i][0:3] == '  -':
            y = doc.add_paragraph()
            g = diff[i][2]
            y.add_run(g)
        elif diff[i][0:3] == '  •':
            y = doc.add_paragraph()
            g = diff[i][2]
            y.add_run(g)
        elif diff[i][0] == '+':
            w = len(diff[i])
            g = diff[i][1:w]
            y.add_run(g).font.color.rgb = RGBColor(255, 0, 0)
        elif diff[i][0] == '-':
            w = len(diff[i])
            g = diff[i][1:w]
            y.add_run(g).font.strike = True
        elif diff[i][0] == '?':
            pass
        else:
            if diff[i] != '+ |':
                y.add_run(diff[i])

    return doc


'''function places text into doc highlighing new and old text'''


def compare_text_newandold(text_1, text_2, doc):
    comp = difflib.Differ()
    diff = list(comp.compare(text_2.split(), text_1.split()))
    new_text = diff
    y = doc.add_paragraph()

    for i in range(0, len(diff)):
        f = len(diff) - 1
        if i < f:
            a = i - 1
        else:
            a = i

        if diff[i][0:3] == '  |':
            j = i + 1
            if diff[i][0:3] and diff[a][0:3] == '  |':
                y = doc.add_paragraph()
            else:
                pass
        elif diff[i][0:3] == '+ |':
            if diff[i][0:3] and diff[a][0:3] == '+ |':
                y = doc.add_paragraph()
            else:
                pass
        elif diff[i][0:3] == '- |':
            pass
        elif diff[i][0:3] == '  -':
            y = doc.add_paragraph()
            g = diff[i][2]
            y.add_run(g)
        elif diff[i][0:3] == '  •':
            y = doc.add_paragraph()
            g = diff[i][2]
            y.add_run(g)
        elif diff[i][0] == '+':
            w = len(diff[i])
            g = diff[i][1:w]
            y.add_run(g).font.color.rgb = RGBColor(255, 0, 0)
        elif diff[i][0] == '-':
            pass
        elif diff[i][0] == '?':
            pass
        else:
            if diff[i] != '+ |':
                y.add_run(diff[i][1:])

    return doc


def printing(dictionary_1, dictionary_2):
    doc = Document()
    new_para = doc.add_paragraph()
    heading = 'Annex B - Dept commentary on actions planned or taken on the IPA RAG rating'
    new_para.add_run(str(heading)).bold = True
    # TODO: change heading font size
    # TODO: be able to change text sixe and font

    for project_name in dictionary_1:
        new_para = doc.add_paragraph()
        new_para.add_run(str(project_name))
        #new_para = doc.add_paragraph()
        dca_text = dictionary_1[project_name]['DCA narrative']
        try:
            dca_text_old = dictionary_2[project_name]['DCA narrative']
        except KeyError:
            dca_text_old = dca_text
        compare_text_newandold(dca_text, dca_text_old, doc)

    new_para = doc.add_paragraph()
    heading = 'Annex C - Narratives against schedule'
    new_para.add_run(str(heading)).bold = True

    for project_name in dictionary_1:
        new_para = doc.add_paragraph()
        new_para.add_run(str(project_name))
        #new_para = doc.add_paragraph()
        dca_text = dictionary_1[project_name]['Dpt narrative on schedule']
        try:
            dca_text_old = dictionary_2[project_name]['Dpt narrative on schedule']
        except KeyError:
            dca_text_old = dca_text
        compare_text_newandold(dca_text, dca_text_old, doc)

    new_para = doc.add_paragraph()
    heading = 'Annex D - departmental narrative on budget/forecast variance of 2018/19 (if variance more than 5%)'
    new_para.add_run(str(heading)).bold = True

    for project_name in dictionary_1:
        new_para = doc.add_paragraph()
        new_para.add_run(str(project_name))
        # new_para = doc.add_paragraph()
        dca_text = dictionary_1[project_name]['narrative on variance']
        try:
            dca_text_old = dictionary_2[project_name]['narrative on variance']
        except KeyError:
            dca_text_old = dca_text
        compare_text_newandold(dca_text, dca_text_old, doc)

    new_para = doc.add_paragraph()
    heading = 'Annex E - narratives against whole life cost (WLC)'
    new_para.add_run(str(heading)).bold = True

    for project_name in dictionary_1:
        new_para = doc.add_paragraph()
        new_para.add_run(str(project_name))
        # new_para = doc.add_paragraph()
        dca_text = dictionary_1[project_name]['narrative on wlc']
        try:
            dca_text_old = dictionary_2[project_name]['narrative on wlc']
        except KeyError:
            dca_text_old = dca_text
        compare_text_newandold(dca_text, dca_text_old, doc)


    return doc


current_Q_dict = project_data_from_master('C:\\Users\\Standalone\\Will\\ipa_annual_report_2019_narratives.xlsx')
last_Q_dict = project_data_from_master('C:\\Users\\Standalone\\Will\\ipa_annual_report_2018.xlsx')

current_Q_list = list(current_Q_dict.keys())

run = printing(current_Q_dict, last_Q_dict)
run.save('C://Users//Standalone//Will//2019_IPA_annual_report_narratives_annex.docx')